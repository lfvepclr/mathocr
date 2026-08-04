"""
Document export module — Markdown and Word (.docx) export.

Markdown export: concatenates all page markdown files.
Word export: parses markdown structure and reconstructs document layout
using python-docx, preserving headers, tables, images, and formulas.
"""

import base64
import html
import json
import logging
import re
import tempfile
import urllib.parse
import urllib.request
from html.parser import HTMLParser
from pathlib import Path

from app.services.batch_manager import (
    BATCHES_DIR,
    get_file_index,
    get_files,
    get_page,
    get_pages,
)
from app.utils.latex import latex_to_unicode

logger = logging.getLogger(__name__)

# Regex patterns
HEADER_RE = re.compile(r"^(#{1,6})\s+(.*)")
IMG_SRC_RE = re.compile(r'src="([^"]+)"')
API_IMG_RE = re.compile(r'/api/page_image/([^/]+)/([^/]+)/(\d+)/([^"]+)')
LATEX_INLINE_RE = re.compile(r"\$([^$]+)\$")
LATEX_BLOCK_RE = re.compile(r"\$\$([^$]+)\$\$", re.DOTALL)
# \(...\) / \[...\] delimiters (siliconflow spotting output uses these)
LATEX_PAREN_RE = re.compile(r"\\\((.+?)\\\)", re.DOTALL)
LATEX_BRACKET_RE = re.compile(r"\\\[(.+?)\\\]", re.DOTALL)

_IMG_MIME_MAP = {
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".png": "image/png",
    ".gif": "image/gif",
    ".bmp": "image/bmp",
    ".webp": "image/webp",
}

_HTTP_IMG_TIMEOUT = 15


def _download_to_temp(url: str) -> Path | None:
    """Best-effort download of a remote image to a temp file.

    Baidu table markdown references figures via bcebos URLs whose auth
    lasts 30 days. New batches localize them at ingest (engine_baidu);
    this fallback keeps exports working for batches created before that
    fix, while the URLs are still valid.
    """
    try:
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=_HTTP_IMG_TIMEOUT) as resp:
            data = resp.read()
        suffix = Path(urllib.parse.urlparse(url).path).suffix.lower()
        if suffix not in _IMG_MIME_MAP:
            suffix = ".png"
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
            tmp.write(data)
        return Path(tmp.name)
    except (OSError, ValueError):
        logger.warning("Export: could not download remote image %s", url[:120])
        return None


# ---------------------------------------------------------------------------
# HTML table parser
# ---------------------------------------------------------------------------
class _TableParser(HTMLParser):
    """Extract tables from HTML. Each cell is a dict with 'text' and 'images'."""

    def __init__(self):
        super().__init__()
        self.tables: list[list[list[dict]]] = []
        self._cur_table: list[list[dict]] | None = None
        self._cur_row: list[dict] | None = None
        self._cur_cell: list[str] | None = None
        self._cur_cell_imgs: list[str] | None = None

    def handle_starttag(self, tag, attrs):
        t = tag.lower()
        if t == "table":
            self._cur_table = []
        elif t == "tr" and self._cur_table is not None:
            self._cur_row = []
        elif t in ("td", "th") and self._cur_row is not None:
            self._cur_cell = []
            self._cur_cell_imgs = []
        elif t == "img" and self._cur_cell is not None and self._cur_cell_imgs is not None:
            for attr_name, attr_val in attrs:
                if attr_name.lower() == "src":
                    self._cur_cell_imgs.append(attr_val)
                    break

    def handle_endtag(self, tag):
        t = tag.lower()
        if t == "table" and self._cur_table is not None:
            self.tables.append(self._cur_table)
            self._cur_table = None
        elif t == "tr" and self._cur_table is not None and self._cur_row is not None:
            self._cur_table.append(self._cur_row)
            self._cur_row = None
        elif t in ("td", "th") and self._cur_row is not None and self._cur_cell is not None:
            self._cur_row.append(
                {
                    "text": "".join(self._cur_cell).strip(),
                    "images": list(self._cur_cell_imgs or []),
                }
            )
            self._cur_cell = None
            self._cur_cell_imgs = None

    def handle_data(self, data):
        if self._cur_cell is not None:
            self._cur_cell.append(data)


def _parse_html_tables(html: str) -> list[list[list[dict]]]:
    """Extract all tables from an HTML string."""
    parser = _TableParser()
    parser.feed(html)
    return parser.tables


# ---------------------------------------------------------------------------
# Image resolution
# ---------------------------------------------------------------------------
def _resolve_image_src(src: str, batch_id: str, file_id=None, page_id=None) -> Path | None:
    """Resolve an image src to a local file path, if possible."""
    # API path: /api/page_image/{batch_id}/{file_id}/{page_id}/{img_name}
    m = API_IMG_RE.match(src)
    if m:
        bid, fid, pid, name = m.groups()
        return BATCHES_DIR / bid / "results" / fid / f"page_{pid}_images" / name

    # Base64 data URI
    if src.startswith("data:image/"):
        try:
            _, data = src.split(",", 1)
            img_bytes = base64.b64decode(data)
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                tmp.write(img_bytes)
            return Path(tmp.name)
        except (ValueError, OSError):
            return None

    # External URL — download at export time (bcebos auth lasts 30 days)
    if src.startswith("http"):
        return _download_to_temp(src)

    # Image filename relative to a page's images dir (baidu block content)
    if file_id is not None and page_id is not None:
        if "/" not in src:
            return BATCHES_DIR / batch_id / "results" / file_id / f"page_{page_id}_images" / src
        # Old local batches prefix table images with "imgs/" though files
        # are stored flat — resolve those by bare file name.
        candidate = (
            BATCHES_DIR / batch_id / "results" / file_id / f"page_{page_id}_images" / Path(src).name
        )
        if candidate.exists():
            return candidate

    # Relative path
    return None


# ---------------------------------------------------------------------------
# Markdown export
# ---------------------------------------------------------------------------
_IMG_SRC_API_RE = re.compile(r'src="/api/page_image/([^/]+)/([^/]+)/(\d+)/([^"]+)"')
_IMG_SRC_HTTP_RE = re.compile(r'src="(https?://[^"]+)"')


def _embed_remote_images_base64(md_text: str) -> str:
    """Best-effort: replace http(s) image URLs with base64 data URIs."""

    def _sub(m: re.Match) -> str:
        tmp = _download_to_temp(m.group(1))
        if not tmp:
            return m.group(0)
        mime = _IMG_MIME_MAP.get(tmp.suffix.lower(), "image/png")
        b64 = base64.b64encode(tmp.read_bytes()).decode("ascii")
        return f'src="data:{mime};base64,{b64}"'

    return _IMG_SRC_HTTP_RE.sub(_sub, md_text)


def _embed_images_base64(md_text: str) -> str:
    """Replace /api/page_image/... URLs with base64 data URIs.

    Exported markdown is meant to be shared outside this platform; API URLs
    break as soon as the file leaves the server. Data URIs make the .md file
    fully self-contained (any markdown viewer renders them inline).
    """

    def _sub(m: re.Match) -> str:
        b_id, f_id, page_id, img_name = m.group(1), m.group(2), m.group(3), m.group(4)
        img_path = BATCHES_DIR / b_id / "results" / f_id / f"page_{page_id}_images" / img_name
        if not img_path.exists():
            logger.warning("Image not found for base64 embedding: %s", img_path)
            return m.group(0)
        mime = _IMG_MIME_MAP.get(img_path.suffix.lower(), "image/png")
        b64 = base64.b64encode(img_path.read_bytes()).decode("ascii")
        return f'src="data:{mime};base64,{b64}"'

    # Remote URLs (bcebos table figures in older baidu batches) are fetched
    # at export time so the .md file stays self-contained
    return _embed_remote_images_base64(_IMG_SRC_API_RE.sub(_sub, md_text))


def export_markdown(batch_id: str, file_id: str) -> str:
    """
    Export a single file's markdown (all pages combined).

    Returns path to the exported .md file.
    """
    pages = get_pages(batch_id, file_id)
    exports_dir = BATCHES_DIR / batch_id / "exports"
    exports_dir.mkdir(parents=True, exist_ok=True)

    # Find original name
    files = get_files(batch_id)
    original_name = next((f["original_name"] for f in files if f["file_id"] == file_id), file_id)
    stem = Path(original_name).stem

    parts = []
    for p in pages:
        md_path = Path(p["markdown_path"]) if p["markdown_path"] else None
        if md_path and md_path.exists():
            parts.append(md_path.read_text(encoding="utf-8"))
            parts.append("\n\n---\n\n")  # page separator

    file_index = get_file_index(batch_id, file_id)
    out_path = exports_dir / f"{batch_id}_{file_index}_{stem}.md"
    out_path.write_text(_embed_images_base64("\n".join(parts)), encoding="utf-8")
    return str(out_path)


def export_batch_markdown(batch_id: str) -> str:
    """Export all files in a batch as a single markdown file.

    Each file is preceded by a visible separator + file-name heading so the
    reader can tell where one document ends and the next begins.
    """
    files = get_files(batch_id)
    exports_dir = BATCHES_DIR / batch_id / "exports"
    exports_dir.mkdir(parents=True, exist_ok=True)

    parts = []
    for i, f in enumerate(files):
        if i > 0:
            parts.append("\n\n---\n\n")  # file separator (horizontal rule)
        parts.append(f"# {f['original_name']}\n\n")
        pages = get_pages(batch_id, f["file_id"])
        for j, p in enumerate(pages):
            md_path = Path(p["markdown_path"]) if p["markdown_path"] else None
            if md_path and md_path.exists():
                if j > 0:
                    parts.append("\n\n---\n\n")  # page separator
                parts.append(md_path.read_text(encoding="utf-8"))

    out_path = exports_dir / f"batch_{batch_id}.md"
    out_path.write_text(_embed_images_base64("\n".join(parts)), encoding="utf-8")
    return str(out_path)


# ---------------------------------------------------------------------------
# Word export
# ---------------------------------------------------------------------------
def export_word(batch_id: str, file_id: str) -> str:
    """
    Export a single file as a Word (.docx) document.

    Parses the markdown structure and reconstructs document layout:
    headers, paragraphs, tables, images, and formulas.
    """
    from docx import Document
    from docx.shared import Pt

    pages = get_pages(batch_id, file_id)
    exports_dir = BATCHES_DIR / batch_id / "exports"
    exports_dir.mkdir(parents=True, exist_ok=True)

    files = get_files(batch_id)
    original_name = next((f["original_name"] for f in files if f["file_id"] == file_id), file_id)
    stem = Path(original_name).stem

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Helvetica"

    for page_idx, page_info in enumerate(pages):
        if page_idx > 0:
            doc.add_page_break()

        md_path = Path(page_info["markdown_path"]) if page_info["markdown_path"] else None
        if not md_path or not md_path.exists():
            doc.add_paragraph(f"[Page {page_info['page_id']} — no content]")
            continue

        md_text = md_path.read_text(encoding="utf-8")
        _add_markdown_to_doc(doc, md_text, batch_id)

    file_index = get_file_index(batch_id, file_id)
    out_path = exports_dir / f"{batch_id}_{file_index}_{stem}.docx"
    doc.save(str(out_path))
    logger.info("Word export saved: %s", out_path)
    return str(out_path)


def export_batch_word(batch_id: str) -> str:
    """Export all files in a batch as a single Word (.docx) document.

    Each file starts on a new page with its name as a title, so documents are
    clearly separated. Files follow the batch's (filename-sorted) order.
    """
    from docx import Document
    from docx.shared import Pt

    files = get_files(batch_id)
    exports_dir = BATCHES_DIR / batch_id / "exports"
    exports_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Helvetica"

    for file_idx, f in enumerate(files):
        if file_idx > 0:
            doc.add_page_break()
        # File separator: original filename as the document title
        doc.add_heading(f["original_name"], level=0)

        pages = get_pages(batch_id, f["file_id"])
        for page_idx, page_info in enumerate(pages):
            if page_idx > 0:
                doc.add_page_break()
            md_path = Path(page_info["markdown_path"]) if page_info["markdown_path"] else None
            if not md_path or not md_path.exists():
                doc.add_paragraph(f"[Page {page_info['page_id']} — no content]")
                continue
            _add_markdown_to_doc(doc, md_path.read_text(encoding="utf-8"), batch_id)

    out_path = exports_dir / f"batch_{batch_id}.docx"
    doc.save(str(out_path))
    logger.info("Batch Word export saved: %s", out_path)
    return str(out_path)


def _add_markdown_to_doc(doc, md_text: str, batch_id: str):
    """Parse markdown text and add elements to a Word document."""

    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        # Skip empty lines
        if not line:
            i += 1
            continue

        # Headers
        m = HEADER_RE.match(line)
        if m:
            hashes, content = m.groups()
            level = len(hashes)
            # Remove markdown formatting from header, convert LaTeX to Unicode
            clean = _convert_latex_in_text(_strip_html_tags(content).strip())
            doc.add_heading(clean, level=min(level, 6))
            i += 1
            continue

        # HTML table
        if line.startswith("<table"):
            table_html = line
            while i + 1 < len(lines) and not lines[i].strip().endswith("</table>"):
                i += 1
                table_html += "\n" + lines[i]
            i += 1
            _add_table_to_doc(doc, table_html, batch_id)
            continue

        # Image (standalone or in div)
        if "<img" in line:
            _add_images_from_html(doc, line, batch_id)
            i += 1
            continue

        # Skip pure div wrappers
        if line.startswith("<div") and "<img" not in line and "<table" not in line:
            i += 1
            continue
        if line.startswith("</div"):
            i += 1
            continue

        # Regular paragraph — collect consecutive non-empty, non-HTML lines
        para_lines = [line]
        while i + 1 < len(lines):
            nxt = lines[i + 1].strip()
            if not nxt or nxt.startswith(("#", "<table", "<div")) or "<img" in nxt:
                break
            para_lines.append(nxt)
            i += 1

        para_text = " ".join(para_lines)
        para_text = _strip_html_tags(para_text).strip()
        if para_text:
            _add_formatted_paragraph(doc, para_text)
        i += 1


def _add_table_to_doc(doc, table_html: str, batch_id: str):
    """Parse HTML table and add to Word document, including images in cells."""
    from docx.shared import Inches

    tables = _parse_html_tables(table_html)
    if not tables:
        return

    table_data = tables[0]
    if not table_data:
        return

    rows = len(table_data)
    cols = max(len(row) for row in table_data) if table_data else 0
    if rows == 0 or cols == 0:
        return

    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"

    for r, row_data in enumerate(table_data):
        for c, cell_data in enumerate(row_data):
            if c >= cols:
                continue
            cell = table.cell(r, c)
            # Extract text (convert LaTeX formulas to Unicode symbols)
            text = _convert_latex_in_text(_strip_html_tags(cell_data.get("text", "")).strip())
            if text:
                cell.text = text
            # Insert images if any
            for src in cell_data.get("images", []):
                img_path = _resolve_image_src(src, batch_id)
                if img_path and img_path.exists():
                    try:
                        paragraph = cell.add_paragraph()
                        run = paragraph.add_run()
                        run.add_picture(str(img_path), width=Inches(1.2))
                    except (OSError, ValueError):
                        logger.debug("Could not insert image in table cell: %s", img_path)


def _add_images_from_html(doc, html_line: str, batch_id: str):
    """Extract and insert images from an HTML line containing <img> tags."""
    from docx.shared import Inches

    for m in IMG_SRC_RE.finditer(html_line):
        src = m.group(1)
        img_path = _resolve_image_src(src, batch_id)
        if img_path and img_path.exists():
            try:
                # Determine width — use original or cap at 5 inches
                from PIL import Image as PILImage

                with PILImage.open(str(img_path)) as pil_img:
                    w, _ = pil_img.size
                # Cap width at 5 inches (assuming 96 DPI for display)
                max_width = min(w / 96, 5.0)
                doc.add_picture(str(img_path), width=Inches(max_width))
                doc.add_paragraph()  # spacing after image
            except (OSError, ValueError):
                logger.debug("Could not insert image: %s", img_path)


def _add_formatted_paragraph(doc, text: str):
    """Add a paragraph with basic markdown formatting (bold, italic, LaTeX)."""

    # Convert LaTeX to Unicode symbols (readable in Word)
    text = _convert_latex_in_text(text)

    # Parse bold and italic
    para = doc.add_paragraph()
    _add_runs_with_formatting(para, text)


def _add_runs_with_formatting(para, text: str):
    """Add text to paragraph with **bold** and *italic* formatting."""
    import re as _re

    # Split by bold markers
    parts = _re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            # Split by italic markers within non-bold parts
            sub_parts = _re.split(r"(\*[^*]+\*)", part)
            for sp in sub_parts:
                if sp.startswith("*") and sp.endswith("*") and len(sp) > 2:
                    run = para.add_run(sp[1:-1])
                    run.italic = True
                else:
                    if sp:
                        para.add_run(sp)


def _convert_latex_in_text(text: str) -> str:
    r"""Convert $...$ / $$...$$ / \(...\) / \[...\] LaTeX segments in text
    to Unicode symbols."""
    text = LATEX_BRACKET_RE.sub(lambda m: f"  {latex_to_unicode(m.group(1).strip())}  ", text)
    text = LATEX_BLOCK_RE.sub(lambda m: f"  {latex_to_unicode(m.group(1).strip())}  ", text)
    text = LATEX_PAREN_RE.sub(lambda m: latex_to_unicode(m.group(1).strip()), text)
    text = LATEX_INLINE_RE.sub(lambda m: latex_to_unicode(m.group(1).strip()), text)
    return text


def _strip_html_tags(text: str) -> str:
    """Remove HTML tags from text, keeping content."""
    return re.sub(r"<[^>]+>", "", text)


# ---------------------------------------------------------------------------
# Word-friendly HTML (copy-to-Word) and layout-faithful HTML export
# ---------------------------------------------------------------------------
_CENTER_DIV_RE = re.compile(r"^<div[^>]*text-align:\s*center[^>]*>(.*)</div>$", re.DOTALL)
_IMG_TAG_RE = re.compile(r"<img[^>]*>")
_BOLD_RE = re.compile(r"\*\*([^*]+)\*\*")
_ITALIC_RE = re.compile(r"\*([^*]+)\*")
_IMAGE_LABELS = ("image", "chart", "seal")


def _img_src_to_data_uri(src: str, batch_id: str, file_id=None, page_id=None) -> str | None:
    img_path = _resolve_image_src(src, batch_id, file_id, page_id)
    if not img_path or not img_path.exists():
        return None
    mime = _IMG_MIME_MAP.get(img_path.suffix.lower(), "image/png")
    b64 = base64.b64encode(img_path.read_bytes()).decode("ascii")
    return f"data:{mime};base64,{b64}"


def _embed_html_images(html_text: str, batch_id: str, file_id=None, page_id=None) -> str:
    """Replace every src=\"...\" in an HTML fragment with a data URI."""

    def _sub(m: re.Match) -> str:
        uri = _img_src_to_data_uri(m.group(1), batch_id, file_id, page_id)
        return f'src="{uri}"' if uri else m.group(0)

    return IMG_SRC_RE.sub(_sub, html_text)


def _inline_text_to_html(text: str) -> str:
    """Escape + LaTeX->Unicode + **bold**/*italic* for inline md text."""
    out = html.escape(_convert_latex_in_text(text))
    out = _BOLD_RE.sub(r"<b>\1</b>", out)
    out = _ITALIC_RE.sub(r"<i>\1</i>", out)
    return out


def _escape_keep_imgs(text: str) -> str:
    """Escape text but keep raw <img> tags intact."""
    out = []
    pos = 0
    for m in _IMG_TAG_RE.finditer(text):
        out.append(_inline_text_to_html(text[pos : m.start()]))
        out.append(m.group(0))
        pos = m.end()
    out.append(_inline_text_to_html(text[pos:]))
    return "".join(out)


def _pipe_table_to_html(md_table: str) -> str | None:
    """Convert a markdown pipe table into an HTML table (Word-pasteable)."""
    rows = []
    for line in md_table.strip().splitlines():
        line = line.strip()
        if not line.startswith("|"):
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        if cells and all(set(c) <= set("-: ") for c in cells if c):
            continue  # separator row
        rows.append(cells)
    if not rows:
        return None
    parts = [
        (
            '<table border="1" cellspacing="0" cellpadding="4" '
            'style="border-collapse:collapse;width:100%">'
        )
    ]
    for r_i, cells in enumerate(rows):
        tag = "th" if r_i == 0 else "td"
        parts.append("<tr>")
        for c in cells:
            parts.append(
                f'<{tag} style="border:1px solid #999;padding:4px">{_escape_keep_imgs(c)}</{tag}>'
            )
        parts.append("</tr>")
    parts.append("</table>")
    return "".join(parts)


def _md_to_word_html(md_text: str, batch_id: str, file_id=None, page_id=None) -> str:
    """Convert page/block markdown into Word-friendly HTML.

    Headings, paragraphs, pipe/HTML tables, centered figures/captions;
    images embedded as data URIs; LaTeX becomes Unicode symbols.
    """
    lines = md_text.split("\n")
    out: list[str] = []
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        m = HEADER_RE.match(line)
        if m:
            hashes, content = m.groups()
            level = min(len(hashes), 6)
            clean = html.escape(_convert_latex_in_text(_strip_html_tags(content).strip()))
            out.append(f"<h{level}>{clean}</h{level}>")
            i += 1
            continue

        if line.startswith("<table"):
            table_html = line
            while "</table>" not in table_html and i + 1 < len(lines):
                i += 1
                table_html += "\n" + lines[i]
            i += 1
            out.append(table_html)
            continue

        # Pipe table (baidu tables)
        if line.startswith("|"):
            tbl_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl_lines.append(lines[i])
                i += 1
            tbl = _pipe_table_to_html("\n".join(tbl_lines))
            if tbl:
                out.append(tbl)
            continue

        # Display formula on its own line
        if line.startswith("$$") and line.endswith("$$") and len(line) > 4:
            body = html.escape(latex_to_unicode(line[2:-2].strip()))
            out.append(f'<p align="center">{body}</p>')
            i += 1
            continue

        # Centered div (figures / captions)
        cm = _CENTER_DIV_RE.match(line)
        if cm:
            out.append(f'<p align="center">{_escape_keep_imgs(cm.group(1).strip())}</p>')
            i += 1
            continue

        # Standalone image line
        if "<img" in line:
            out.append(f'<p align="center">{_escape_keep_imgs(line)}</p>')
            i += 1
            continue

        # Skip pure div wrappers
        if line.startswith(("<div", "</div")):
            i += 1
            continue

        # Paragraph — collect consecutive plain lines
        para_lines = [line]
        while i + 1 < len(lines):
            nxt = lines[i + 1].strip()
            if not nxt or nxt.startswith(("#", "|", "<table", "<div")) or "<img" in nxt:
                break
            para_lines.append(nxt)
            i += 1
        out.append(f"<p>{_escape_keep_imgs(' '.join(para_lines))}</p>")
        i += 1

    return _embed_html_images("\n".join(out), batch_id, file_id, page_id)


def _block_fragment(blocks: list, idx: int, md_text: str) -> str:
    """Markdown fragment for a single block (image tags and title levels
    restored, so the pasted fragment stands alone in Word)."""
    block = blocks[idx]
    label = block.get("block_label") or ""
    fragment = block.get("block_content") or ""
    if "<img" not in fragment and label in _IMAGE_LABELS:
        # Image blocks carry empty content in local results — the <img>
        # tags live in the full-page markdown, in reading order
        img_tags = _IMG_TAG_RE.findall(md_text)
        order = (
            sum(1 for b in blocks[: idx + 1] if (b.get("block_label") or "") in _IMAGE_LABELS) - 1
        )
        if 0 <= order < len(img_tags):
            fragment = img_tags[order]
    # Titles carry no '#' prefix in block_content — add one so the
    # pasted fragment becomes a real Word heading
    if fragment and not fragment.startswith("#"):
        if label == "doc_title":
            fragment = f"# {fragment}"
        elif label == "paragraph_title":
            fragment = f"## {fragment}"
    return fragment


def page_to_richtext(
    batch_id: str,
    file_id: str,
    page_id: int,
    block_idx: int | None = None,
    block_idxs: list[int] | None = None,
) -> dict:
    """Word-friendly HTML + plain text for a page, a single block, or a
    multi-block selection (lasso copy).

    Powers the frontend's copy-to-Word: the HTML goes to the clipboard as
    text/html so pasting into Word keeps headings/tables/images.
    """
    page = get_page(batch_id, file_id, page_id)
    if not page:
        raise ValueError(f"页面不存在: {batch_id}/{file_id}/{page_id}")
    md_path = Path(page["markdown_path"]) if page["markdown_path"] else None
    md_text = md_path.read_text(encoding="utf-8") if md_path and md_path.exists() else ""

    if block_idxs is not None and not block_idxs:
        raise ValueError("块索引列表为空")
    if block_idx is None and block_idxs is None:
        html_text = _md_to_word_html(md_text, batch_id, file_id, page_id)
        plain = html.unescape(re.sub(r"<[^>]+>", "", html_text))
        plain = re.sub(r"\n{3,}", "\n\n", plain).strip()
        return {"html": html_text, "text": plain}

    json_path = Path(page["json_path"]) if page["json_path"] else None
    blocks = []
    if json_path and json_path.exists():
        data = json.loads(json_path.read_text(encoding="utf-8"))
        blocks = (data.get("res") or data).get("parsing_res_list") or []

    if block_idxs:
        # Multi-block selection: dedupe, then restore reading order
        idxs = []
        for i in block_idxs:
            if not 0 <= i < len(blocks):
                raise ValueError(f"块索引超出范围: {i}")
            if i not in idxs:
                idxs.append(i)
        idxs.sort(key=lambda i: blocks[i].get("block_order") or i)
    else:
        assert block_idx is not None
        if not 0 <= block_idx < len(blocks):
            raise ValueError(f"块索引超出范围: {block_idx}")
        idxs = [block_idx]

    html_parts, text_parts = [], []
    for i in idxs:
        fragment = _block_fragment(blocks, i, md_text)
        part = _md_to_word_html(fragment, batch_id, file_id, page_id)
        html_parts.append(part)
        text = html.unescape(re.sub(r"<[^>]+>", "", part)).strip()
        if text:
            text_parts.append(text)
    return {"html": "\n".join(html_parts), "text": "\n\n".join(text_parts)}


# ---------------------------------------------------------------------------
# Layout-faithful HTML export
# ---------------------------------------------------------------------------
_LAYOUT_TEMPLATE = """<!DOCTYPE html>
<html lang="zh-CN"><head><meta charset="utf-8">
<title>__TITLE__</title>
<style>
body{margin:0;background:#e5e7eb;font-family:"Songti SC","SimSun",serif;}
.page{position:relative;width:900px;margin:16px auto;background:#fff;
  box-shadow:0 1px 6px rgba(0,0,0,.25);overflow:hidden;}
.lb{position:absolute;overflow:hidden;}
.lb .t{font-weight:700;}
.lb .t-doc_title{text-align:center;}
.lb .f{text-align:center;}
.lb-img img{width:100%;height:100%;object-fit:fill;display:block;}
.lb-tbl table{width:100%;height:100%;border-collapse:collapse;table-layout:fixed;}
.lb-tbl td,.lb-tbl th{border:1px solid #666;padding:2px;
  vertical-align:middle;text-align:center;overflow:hidden;}
.lb-tbl img{max-width:100%;max-height:100%;}
/* Per-label facsimile styling (mirrors the web layout view) */
.lb[data-label="paragraph_title"]{font-weight:700;}
.lb[data-label="figure_title"],.lb[data-label="table_title"]{
  font-size:.85em;text-align:center;}
.lb[data-label="header"],.lb[data-label="footer"],.lb[data-label="number"]{
  color:#64748b;font-size:.85em;}
.lb[data-label="number"]{text-align:center;}
.lb[data-label="formula"]{text-align:center;}
.lb[data-label="formula_number"]{text-align:right;}
.lb[data-label="vertical_text"]{writing-mode:vertical-rl;}
.lb[data-label="footnote"],.lb[data-label="reference_content"]{font-size:.85em;}
@media print{
  body{background:#fff;}
  .page{margin:0 auto;box-shadow:none;page-break-after:always;}
}
</style></head><body>
__PAGES__
<script>
// Shrink-to-fit: estimate font size from box geometry, then shrink until
// the content fits its absolutely-positioned box (min 6px, then clip).
document.querySelectorAll('.lb').forEach(function(el){
  if (el.classList.contains('lb-img')) return;
  var W = el.clientWidth, H = el.clientHeight;
  if (!W || !H) return;
  var L = el.textContent.length || 1;
  var fs = Math.min(H * 0.8, Math.sqrt(W * H * 1.9 / L));
  el.style.fontSize = fs + 'px';
  el.style.lineHeight = '1.25';
  for (var i = 0; i < 30; i++) {
    if (el.scrollHeight <= el.clientHeight + 1 &&
        el.scrollWidth <= el.clientWidth + 1) break;
    fs *= 0.92;
    el.style.fontSize = fs + 'px';
    if (fs < 6) break;
  }
});
</script></body></html>
"""


def _layout_page_html(batch_id: str, file_id: str, page_info: dict) -> str:
    """One absolutely-positioned page canvas for the layout HTML export."""
    json_path = Path(page_info["json_path"]) if page_info["json_path"] else None
    if not json_path or not json_path.exists():
        return ""
    data = json.loads(json_path.read_text(encoding="utf-8"))
    res = data.get("res") or data
    width = res.get("width") or 0
    height = res.get("height") or 0
    blocks = res.get("parsing_res_list") or []
    if not width or not height:
        return ""

    md_path = Path(page_info["markdown_path"]) if page_info["markdown_path"] else None
    md_text = md_path.read_text(encoding="utf-8") if md_path and md_path.exists() else ""
    # Image blocks in local results have empty content — <img> tags live in
    # the full-page markdown, in reading order
    img_tags = _IMG_TAG_RE.findall(md_text)
    img_cursor = 0
    page_id = page_info["page_id"]

    parts = [f'<div class="page" style="aspect-ratio:{width}/{height}">']
    for block in blocks:
        bbox = block.get("block_bbox")
        if not bbox or len(bbox) != 4:
            continue
        x1, y1, x2, y2 = bbox
        if x2 <= x1 or y2 <= y1:
            continue
        label = block.get("block_label") or "text"
        content = block.get("block_content") or ""

        if label in _IMAGE_LABELS:
            tag = ""
            m = _IMG_TAG_RE.search(content)
            if m:
                tag = m.group(0)
            elif img_cursor < len(img_tags):
                tag = img_tags[img_cursor]
                img_cursor += 1
            uri = None
            if tag:
                src_m = IMG_SRC_RE.search(tag)
                if src_m:
                    uri = _img_src_to_data_uri(src_m.group(1), batch_id, file_id, page_id)
            if not uri:
                continue
            body = f'<img src="{uri}" alt="">'
            cls = "lb lb-img"
        elif label == "table":
            if content.lstrip().startswith("|"):
                tbl = _pipe_table_to_html(content)
            elif "<table" in content:
                tbl = content
            else:
                tbl = None
            if tbl is None:
                tbl = f"<p>{_escape_keep_imgs(content)}</p>"
            body = _embed_html_images(tbl, batch_id, file_id, page_id)
            cls = "lb lb-tbl"
        elif label == "formula":
            txt = content.strip()
            if txt.startswith("$$") and txt.endswith("$$"):
                txt = txt[2:-2].strip()
            elif "$" not in txt:
                txt = f"${txt}$"
            body = f'<div class="f">{html.escape(_convert_latex_in_text(txt))}</div>'
            cls = "lb"
        elif label in ("doc_title", "paragraph_title"):
            body = (
                f'<div class="t t-{label}">{_inline_text_to_html(_strip_html_tags(content))}</div>'
            )
            cls = "lb"
        else:
            body = f'<div class="x">{_escape_keep_imgs(content)}</div>'
            cls = "lb"

        left = x1 / width * 100
        top = y1 / height * 100
        w_pct = (x2 - x1) / width * 100
        h_pct = (y2 - y1) / height * 100
        parts.append(
            f'<div class="{cls}" data-label="{html.escape(label)}" '
            f'style="left:{left:.3f}%;top:{top:.3f}%;'
            f'width:{w_pct:.3f}%;height:{h_pct:.3f}%">{body}</div>'
        )
    parts.append("</div>")
    return "".join(parts)


def export_layout_html(batch_id: str, file_id: str | None = None) -> str:
    """Self-contained HTML reproducing the original page layout.

    Every block is absolutely positioned at its bbox on a page canvas;
    images are embedded as data URIs; an inline script shrink-fits each
    block's font. Printable via the browser (one .page per printed page).
    """
    files = get_files(batch_id)
    if file_id:
        files = [f for f in files if f["file_id"] == file_id]
    exports_dir = BATCHES_DIR / batch_id / "exports"
    exports_dir.mkdir(parents=True, exist_ok=True)

    pages_html = []
    for f in files:
        for p in get_pages(batch_id, f["file_id"]):
            pages_html.append(_layout_page_html(batch_id, f["file_id"], p))

    if file_id:
        stem = Path(files[0]["original_name"]).stem
        file_index = get_file_index(batch_id, file_id)
        name = f"{batch_id}_{file_index}_{stem}_layout.html"
        title = f"{stem} — 版面还原"
    else:
        name = f"batch_{batch_id}_layout.html"
        title = f"批次 {batch_id} — 版面还原"

    doc = _LAYOUT_TEMPLATE.replace("__TITLE__", html.escape(title))
    doc = doc.replace("__PAGES__", "\n".join(pages_html))
    out_path = exports_dir / name
    out_path.write_text(doc, encoding="utf-8")
    logger.info("Layout HTML export saved: %s", out_path)
    return str(out_path)
