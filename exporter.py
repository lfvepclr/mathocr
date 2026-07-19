"""
Document export module — Markdown and Word (.docx) export.

Markdown export: concatenates all page markdown files.
Word export: parses markdown structure and reconstructs document layout
using python-docx, preserving headers, tables, images, and formulas.
"""

import base64
import io
import logging
import re
import tempfile
from html.parser import HTMLParser
from pathlib import Path

from batch_manager import (
    BATCHES_DIR,
    get_file_index,
    get_file_results_dir,
    get_files,
    get_pages,
)
from latex_utils import latex_to_unicode

logger = logging.getLogger(__name__)

# Regex patterns
HEADER_RE = re.compile(r"^(#{1,6})\s+(.*)")
IMG_SRC_RE = re.compile(r'src="([^"]+)"')
API_IMG_RE = re.compile(r'/api/page_image/([^/]+)/([^/]+)/(\d+)/([^"]+)')
LATEX_INLINE_RE = re.compile(r"\$([^$]+)\$")
LATEX_BLOCK_RE = re.compile(r"\$\$([^$]+)\$\$", re.DOTALL)


# ---------------------------------------------------------------------------
# HTML table parser
# ---------------------------------------------------------------------------
class _TableParser(HTMLParser):
    """Extract tables from HTML. Each cell is a dict with 'text' and 'images'."""

    def __init__(self):
        super().__init__()
        self.tables: list[list[list[dict]]] = []
        self._cur_table: list[list[dict]] | None = None
        self._cur_row: list[dict] | None = None
        self._cur_cell: list[str] | None = None
        self._cur_cell_imgs: list[str] | None = None

    def handle_starttag(self, tag, attrs):
        t = tag.lower()
        if t == "table":
            self._cur_table = []
        elif t == "tr" and self._cur_table is not None:
            self._cur_row = []
        elif t in ("td", "th") and self._cur_row is not None:
            self._cur_cell = []
            self._cur_cell_imgs = []
        elif t == "img" and self._cur_cell is not None:
            for attr_name, attr_val in attrs:
                if attr_name.lower() == "src":
                    self._cur_cell_imgs.append(attr_val)
                    break

    def handle_endtag(self, tag):
        t = tag.lower()
        if t == "table" and self._cur_table is not None:
            self.tables.append(self._cur_table)
            self._cur_table = None
        elif t == "tr" and self._cur_row is not None:
            self._cur_table.append(self._cur_row)
            self._cur_row = None
        elif t in ("td", "th") and self._cur_cell is not None:
            self._cur_row.append({
                "text": "".join(self._cur_cell).strip(),
                "images": list(self._cur_cell_imgs),
            })
            self._cur_cell = None
            self._cur_cell_imgs = None

    def handle_data(self, data):
        if self._cur_cell is not None:
            self._cur_cell.append(data)


def _parse_html_tables(html: str) -> list[list[list[dict]]]:
    """Extract all tables from an HTML string."""
    parser = _TableParser()
    parser.feed(html)
    return parser.tables


# ---------------------------------------------------------------------------
# Image resolution
# ---------------------------------------------------------------------------
def _resolve_image_src(src: str, batch_id: str) -> Path | None:
    """Resolve an image src to a local file path, if possible."""
    # API path: /api/page_image/{batch_id}/{file_id}/{page_id}/{img_name}
    m = API_IMG_RE.match(src)
    if m:
        bid, fid, pid, name = m.groups()
        return (
            BATCHES_DIR / bid / "results" / fid
            / f"page_{pid}_images" / name
        )

    # Base64 data URI
    if src.startswith("data:image/"):
        try:
            header, data = src.split(",", 1)
            img_bytes = base64.b64decode(data)
            tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
            tmp.write(img_bytes)
            tmp.close()
            return Path(tmp.name)
        except Exception:
            return None

    # External URL — cannot embed in Word
    if src.startswith("http"):
        return None

    # Relative path
    return None


# ---------------------------------------------------------------------------
# Markdown export
# ---------------------------------------------------------------------------
_IMG_MIME_MAP = {
    ".jpg": "image/jpeg", ".jpeg": "image/jpeg", ".png": "image/png",
    ".gif": "image/gif", ".bmp": "image/bmp", ".webp": "image/webp",
}

_IMG_SRC_API_RE = re.compile(r'src="/api/page_image/([^/]+)/([^/]+)/(\d+)/([^"]+)"')


def _embed_images_base64(md_text: str) -> str:
    """Replace /api/page_image/... URLs with base64 data URIs.

    Exported markdown is meant to be shared outside this platform; API URLs
    break as soon as the file leaves the server. Data URIs make the .md file
    fully self-contained (any markdown viewer renders them inline).
    """
    def _sub(m: re.Match) -> str:
        b_id, f_id, page_id, img_name = m.group(1), m.group(2), m.group(3), m.group(4)
        img_path = (BATCHES_DIR / b_id / "results" / f_id
                    / f"page_{page_id}_images" / img_name)
        if not img_path.exists():
            logger.warning("Image not found for base64 embedding: %s", img_path)
            return m.group(0)
        mime = _IMG_MIME_MAP.get(img_path.suffix.lower(), "image/png")
        b64 = base64.b64encode(img_path.read_bytes()).decode("ascii")
        return f'src="data:{mime};base64,{b64}"'

    return _IMG_SRC_API_RE.sub(_sub, md_text)


def export_markdown(batch_id: str, file_id: str) -> str:
    """
    Export a single file's markdown (all pages combined).

    Returns path to the exported .md file.
    """
    pages = get_pages(batch_id, file_id)
    exports_dir = BATCHES_DIR / batch_id / "exports"
    exports_dir.mkdir(parents=True, exist_ok=True)

    # Find original name
    files = get_files(batch_id)
    original_name = next(
        (f["original_name"] for f in files if f["file_id"] == file_id), file_id
    )
    stem = Path(original_name).stem

    parts = []
    for p in pages:
        md_path = Path(p["markdown_path"]) if p["markdown_path"] else None
        if md_path and md_path.exists():
            parts.append(md_path.read_text(encoding="utf-8"))
            parts.append("\n\n---\n\n")  # page separator

    file_index = get_file_index(batch_id, file_id)
    out_path = exports_dir / f"{batch_id}_{file_index}_{stem}.md"
    out_path.write_text(_embed_images_base64("\n".join(parts)), encoding="utf-8")
    return str(out_path)


def export_batch_markdown(batch_id: str) -> str:
    """Export all files in a batch as a single markdown file."""
    files = get_files(batch_id)
    exports_dir = BATCHES_DIR / batch_id / "exports"
    exports_dir.mkdir(parents=True, exist_ok=True)

    parts = []
    for f in files:
        parts.append(f"# {f['original_name']}\n\n")
        pages = get_pages(batch_id, f["file_id"])
        for p in pages:
            md_path = Path(p["markdown_path"]) if p["markdown_path"] else None
            if md_path and md_path.exists():
                parts.append(md_path.read_text(encoding="utf-8"))
                parts.append("\n\n---\n\n")

    out_path = exports_dir / f"batch_{batch_id}.md"
    out_path.write_text(_embed_images_base64("\n".join(parts)), encoding="utf-8")
    return str(out_path)


# ---------------------------------------------------------------------------
# Word export
# ---------------------------------------------------------------------------
def export_word(batch_id: str, file_id: str) -> str:
    """
    Export a single file as a Word (.docx) document.

    Parses the markdown structure and reconstructs document layout:
    headers, paragraphs, tables, images, and formulas.
    """
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    pages = get_pages(batch_id, file_id)
    exports_dir = BATCHES_DIR / batch_id / "exports"
    exports_dir.mkdir(parents=True, exist_ok=True)

    files = get_files(batch_id)
    original_name = next(
        (f["original_name"] for f in files if f["file_id"] == file_id), file_id
    )
    stem = Path(original_name).stem

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Helvetica"

    for page_idx, page_info in enumerate(pages):
        if page_idx > 0:
            doc.add_page_break()

        md_path = Path(page_info["markdown_path"]) if page_info["markdown_path"] else None
        if not md_path or not md_path.exists():
            doc.add_paragraph(f"[Page {page_info['page_id']} — no content]")
            continue

        md_text = md_path.read_text(encoding="utf-8")
        _add_markdown_to_doc(doc, md_text, batch_id)

    file_index = get_file_index(batch_id, file_id)
    out_path = exports_dir / f"{batch_id}_{file_index}_{stem}.docx"
    doc.save(str(out_path))
    logger.info("Word export saved: %s", out_path)
    return str(out_path)


def _add_markdown_to_doc(doc, md_text: str, batch_id: str):
    """Parse markdown text and add elements to a Word document."""
    from docx.shared import Inches

    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        # Skip empty lines
        if not line:
            i += 1
            continue

        # Headers
        m = HEADER_RE.match(line)
        if m:
            hashes, content = m.groups()
            level = len(hashes)
            # Remove markdown formatting from header, convert LaTeX to Unicode
            clean = _convert_latex_in_text(_strip_html_tags(content).strip())
            doc.add_heading(clean, level=min(level, 6))
            i += 1
            continue

        # HTML table
        if line.startswith("<table"):
            table_html = line
            while i + 1 < len(lines) and not lines[i].strip().endswith("</table>"):
                i += 1
                table_html += "\n" + lines[i]
            i += 1
            _add_table_to_doc(doc, table_html, batch_id)
            continue

        # Image (standalone or in div)
        if "<img" in line:
            _add_images_from_html(doc, line, batch_id)
            i += 1
            continue

        # Skip pure div wrappers
        if line.startswith("<div") and "<img" not in line and "<table" not in line:
            i += 1
            continue
        if line.startswith("</div"):
            i += 1
            continue

        # Regular paragraph — collect consecutive non-empty, non-HTML lines
        para_lines = [line]
        while i + 1 < len(lines):
            nxt = lines[i + 1].strip()
            if not nxt or nxt.startswith("#") or nxt.startswith("<table") or "<img" in nxt or nxt.startswith("<div"):
                break
            para_lines.append(nxt)
            i += 1

        para_text = " ".join(para_lines)
        para_text = _strip_html_tags(para_text).strip()
        if para_text:
            _add_formatted_paragraph(doc, para_text)
        i += 1


def _add_table_to_doc(doc, table_html: str, batch_id: str):
    """Parse HTML table and add to Word document, including images in cells."""
    from docx.shared import Inches

    tables = _parse_html_tables(table_html)
    if not tables:
        return

    table_data = tables[0]
    if not table_data:
        return

    rows = len(table_data)
    cols = max(len(row) for row in table_data) if table_data else 0
    if rows == 0 or cols == 0:
        return

    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"

    for r, row_data in enumerate(table_data):
        for c, cell_data in enumerate(row_data):
            if c >= cols:
                continue
            cell = table.cell(r, c)
            # Extract text (convert LaTeX formulas to Unicode symbols)
            text = _convert_latex_in_text(_strip_html_tags(cell_data.get("text", "")).strip())
            if text:
                cell.text = text
            # Insert images if any
            for src in cell_data.get("images", []):
                img_path = _resolve_image_src(src, batch_id)
                if img_path and img_path.exists():
                    try:
                        paragraph = cell.add_paragraph()
                        run = paragraph.add_run()
                        run.add_picture(str(img_path), width=Inches(1.2))
                    except Exception:
                        logger.debug("Could not insert image in table cell: %s", img_path)


def _add_images_from_html(doc, html_line: str, batch_id: str):
    """Extract and insert images from an HTML line containing <img> tags."""
    from docx.shared import Inches

    for m in IMG_SRC_RE.finditer(html_line):
        src = m.group(1)
        img_path = _resolve_image_src(src, batch_id)
        if img_path and img_path.exists():
            try:
                # Determine width — use original or cap at 5 inches
                from PIL import Image as PILImage

                with PILImage.open(str(img_path)) as pil_img:
                    w, h = pil_img.size
                # Cap width at 5 inches (assuming 96 DPI for display)
                max_width = min(w / 96, 5.0)
                doc.add_picture(str(img_path), width=Inches(max_width))
                doc.add_paragraph()  # spacing after image
            except Exception:
                logger.debug("Could not insert image: %s", img_path)


def _add_formatted_paragraph(doc, text: str):
    """Add a paragraph with basic markdown formatting (bold, italic, LaTeX)."""
    from docx.shared import Pt

    # Convert LaTeX to Unicode symbols (readable in Word)
    text = _convert_latex_in_text(text)

    # Parse bold and italic
    para = doc.add_paragraph()
    _add_runs_with_formatting(para, text)


def _add_runs_with_formatting(para, text: str):
    """Add text to paragraph with **bold** and *italic* formatting."""
    from docx.shared import Pt
    import re as _re

    # Split by bold markers
    parts = _re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            # Split by italic markers within non-bold parts
            sub_parts = _re.split(r"(\*[^*]+\*)", part)
            for sp in sub_parts:
                if sp.startswith("*") and sp.endswith("*") and len(sp) > 2:
                    run = para.add_run(sp[1:-1])
                    run.italic = True
                else:
                    if sp:
                        para.add_run(sp)


def _convert_latex_in_text(text: str) -> str:
    """Convert $...$ / $$...$$ LaTeX segments in text to Unicode symbols."""
    text = LATEX_BLOCK_RE.sub(lambda m: f"  {latex_to_unicode(m.group(1).strip())}  ", text)
    text = LATEX_INLINE_RE.sub(lambda m: latex_to_unicode(m.group(1).strip()), text)
    return text


def _strip_html_tags(text: str) -> str:
    """Remove HTML tags from text, keeping content."""
    return re.sub(r"<[^>]+>", "", text)
